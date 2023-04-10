import random
import re
import zipfile

from docx import Document
from config import Config
from database import Database

class DocTemplate:
    def __init__(self, file_path: str):
        self.document = Document(file_path)

    def get_patterns_from_template(self):
        pattern_list = []
        paragraphs = []

        for paragraph in self.document.paragraphs:
            paragraphs.append(paragraph)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)

        for paragraph in paragraphs:
            patterns = re.findall("\{\{.+\}\}", paragraph.text)
            for pattern in patterns:
                pattern_data = pattern[3:-3:].split(":")
                if len(pattern_data) == 2:
                    pattern_name, pattern_desc = pattern_data
                else:
                    pattern_name, pattern_desc = pattern_data[0], "Nonei"

                if any(map(lambda i: i["name"] == pattern_name, pattern_list)):
                    continue

                pattern_list.append({
                    "name": pattern_name,
                    "desc": pattern_desc
                })
                
        return pattern_list

    def fill_template(self, data: dict):
        for row in self.document.tables[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    patterns = re.findall("\{\{.+\}\}", paragraph.text)

                    for pattern in patterns:
                        pattern_name = pattern[3:-3:].split(":")[0]
                        
                        if pattern_name in data:
                            paragraph.text = paragraph.text.replace(pattern, data[pattern_name])

        for paragraph in self.document.paragraphs:
            patterns = re.findall("\{\{.+\}\}", paragraph.text)

            for pattern in patterns:
                pattern_name = pattern[3:-3:].split(":")[0]

                if pattern_name in data:
                    self.replace_text_in_paragraph(paragraph, data[pattern_name])

        path = f"temp/{random.randint(0, 999999999)}.docx"
        self.document.save(path)

        return path

    @staticmethod
    def replace_text_in_paragraph(paragraph, value):
        items = []
        close_char_count = 0
        is_open = False

        for item in paragraph.runs:
            if is_open:
                items.append(item)
                close_char_count += item.text.count("}")
            else:
                if "{" in item.text:
                    is_open = True
                    items.append(item)

                    close_char_count += item.text.count("}")

            if close_char_count == 2:
                break

        items[0].text = value

        for item in items[1:]:
            item.text = ""    

    @staticmethod
    def multi_fill_template(
        config: Config, 
        template_id: int, 
        database: str, 
        table: str, 
        data: dict
    ) -> str:
        template_info = config.get_template_info(template_id)
        file_paths = []
        db = Database(database_filename=database)
        table_data = db.get_table_rows(table)

        for row in table_data:
            template = DocTemplate(template_info["file_path"])
            data_for_template = DocTemplate.get_data_for_template(data, row)
            path = template.fill_template(data_for_template)

            file_paths.append(path)

        zip_file_path = f"temp/{random.randint(0, 9**9)}.zip"

        with zipfile.ZipFile(zip_file_path, 'w') as zip:
            for file_path in file_paths:
                zip.write(file_path)

        return zip_file_path

    @staticmethod
    def get_data_for_template(req_data: dict, row: dict):
        data = {}

        for key, value in req_data.items():
            if key in ["id", "database", "table"]:
                continue

            data.update({key: row[value]})

        return data





    